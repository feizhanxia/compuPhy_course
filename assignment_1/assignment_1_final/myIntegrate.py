# how to caculate integration
import numpy as np
import sys
from numba import njit
import time
import matplotlib.pyplot as plt
from matplotlib import ticker
from docx import Document
from docx.shared import Inches

#  def quad(f, a, b, n):  # function of integrate
#  h = (b - a) / n  # step length
#  sum = 0  # initial integration result = 0
#  for ii in range(0, n):
#  sum = (sum + f(a + (ii + 0.5) * h)) * h
#  return sum

#  def adaptive_quad(f, a, b, eps):
#  it = 0  # initial iteration count = 0
#  maxIt = 100  # maximum iteration times = 100
#  error = sys.float_info.max
#  n = 1
#  vOld = quad(f, a, b, n)
#  while error > eps and maxIt > it:
#  n = n * 2
#  vNew = quad(f, a, b, n)
#  error = np.abs(vNew - vOld)
#  vOld = vNew
#  it += 1
#  print("Iteration:", it)
#  return vOld

#  def simpson(f, a, b, n):
#  h = (b - a) / n
#  sum = 0
#  for ii in range(n):
#  sum += 2 * f(a + ii * h) + 4 * f(a + ii * h + 0.5 * h)
#  sum = h * (sum + f(b) - f(a)) / 6
#  return sum

#  def adaptive_simpson(f, a, b, eps):
#  it = 0  # initial iteration count = 0
#  maxIt = 100  # maximum iteration times = 100
#  error = sys.float_info.max
#  n = 1
#  vOld = simpson(f, a, b, n)
#  while error > eps and maxIt > it:
#  n = n * 2
#  vNew = quad(f, a, b, n)
#  error = np.abs(vNew - vOld)
#  vOld = vNew
#  it += 1
#  print("Iteration:", it)
#  return vOld

# 下面进行面向对象的编程 integration


class myIntegrate():
    def __init__(self, func, a, b):
        self.f = func
        self.a = a
        self.b = b
        self.sum = 0
        self.integrateValue = []
        self.integrateError = []
        self.time = 0
        self.method = ''

    def adaptive(self, eps=1e-3, n=1):
        t0 = time.time()
        it = 0  # initial iteration count = 0
        maxIt = 100  # maximum iteration times = 100
        error = sys.float_info.max
        vOld = self.caculate(n)
        self.integrateValue = []
        self.integrateError = []
        while error > eps and maxIt > it:
            n = n * 2
            vNew = self.caculate(n)
            error = np.abs(vNew - vOld)
            vOld = vNew
            self.integrateValue.append(vOld)
            self.integrateError.append(error)
            it += 1
        t1 = time.time()
        self.time = t1 - t0
        print("Iteration:", it)
        return vOld


class quad(myIntegrate):
    def __init__(self, func, a, b):
        super().__init__(func, a, b)
        self.method = 'quad'

    def caculate(self, n):
        points = np.linspace(self.a, self.b, n + 1)[:-1]
        weight = np.full_like(points, (self.b - self.a) / n)  # step length
        self.sum = self.f(points + 0.5 * weight) @ weight
        #  sum = 0  # initial integration result = 0
        #  for ii in range(n):
        #  sum += self.f(self.a + (ii + 0.5) * h)
        #  self.sum = sum * h
        return self.sum


class trapezoidal(myIntegrate):
    def __init__(self, func, a, b):
        super().__init__(func, a, b)
        self.method = 'trapezoidal'

    def caculate(self, n):
        points = np.linspace(self.a, self.b, n + 1)[:-1]
        weight = np.full_like(points, (self.b - self.a) / n)
        self.sum = self.f(points) @ weight + (self.f(self.b) -
                                              self.f(self.a)) * weight[0] / 2
        #  for ii in range(n):
        #  sum = sum + 2 * self.f(self.a + ii * h)
        #  print(sum)
        #  self.sum = h * (sum + self.f(self.b) - self.f(self.a)) / 2
        return self.sum


class simpson(myIntegrate):
    def __init__(self, func, a, b):
        super().__init__(func, a, b)
        self.method = 'simpson'

    def caculate(self, n):
        points = np.linspace(self.a, self.b, n + 1).astype(np.float64)
        points = points[:-1]
        weight = np.full_like(points, (self.b - self.a) / n)
        self.sum = (
            (2 * self.f(points) + 4 * self.f(points + weight * 0.5)) @ weight +
            (self.f(self.b) - self.f(self.a)) * weight[0]) / 6
        #  h = (self.b - self.a) / n
        #  sum = 0
        #  for ii in range(n):
        #  sum += 2 * self.f(self.a + ii * h) + 4 * self.f(self.a + ii * h +
        #  0.5 * h)
        #  self.sum = h * (sum + self.f(self.b) - self.f(self.a)) / 6
        return self.sum


# 4
# 5
class gauss(myIntegrate):
    def __init__(self, func, a, b, m=1):
        super().__init__(func, a, b)
        self.nodeNum = m + 1  # 当 m=1 即default时，为普通高斯积分，当m > 1时为均匀分m段的分段高斯积分
        if m == 1:
            self.method = 'gaussian'
        else:
            self.method = 'gaussian in ' + str(m) + ' parts'

    def partialCaculate(self, a, b, n):
        points, weight = np.polynomial.legendre.leggauss(n)
        partialSum = (b - a) / 2 * self.f((b - a) / 2 * points +
                                          (b + a) / 2) @ weight
        return partialSum

    def caculate(self, n):
        ufuncPartialCaculate = np.frompyfunc(self.partialCaculate, 3, 1)
        nodes = np.linspace(self.a, self.b, self.nodeNum)
        self.sum = ufuncPartialCaculate(nodes[:-1], nodes[1:], n).sum()
        return self.sum


if __name__ == '__main__':

    @njit
    def fun(x):
        return (x**2 + np.cos(x + 2))

    errorGoal = 1e-5
    integralFun1 = trapezoidal(fun, -2, 3)
    res1 = integralFun1.adaptive(errorGoal)
    integralFun2 = simpson(fun, -2, 3)
    res2 = integralFun2.adaptive(errorGoal)
    integralFun3 = gauss(fun, -2, 3)
    res3 = integralFun3.adaptive(errorGoal)
    integralFun4 = gauss(fun, -2, 3, 2)
    res4 = integralFun4.adaptive(errorGoal)

    #  print("Integral is: ", res)
    #  print("error:", integralFun.integrateError[-1])

    # 下面输出word格式的收敛表格
    document = Document()

    for integralFun in [
            integralFun1, integralFun2, integralFun3, integralFun4
    ]:
        table = document.add_table(1, 3)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Iteration'
        heading_cells[1].text = 'Result'
        heading_cells[2].text = 'Error'
        for i, val in enumerate(np.array(integralFun.integrateValue)):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = '%.9g' % val
            cells[2].text = '%.4e' % integralFun.integrateError[i]
        table.style = 'Light Shading Accent 1'
        document.add_paragraph('Integration method: ' + integralFun.method +
                               '. Total caculation time is: ' +
                               ('%.6g' % integralFun.time) +
                               ' , the result is: ' +
                               ('%.9g' % integralFun.sum) + ' .')

    # 下面开始画图
    fig, axs = plt.subplots(1, 2, figsize=(15, 5))
    for integralFun in [
            integralFun1, integralFun2, integralFun3, integralFun4
    ]:
        X = range(len(integralFun.integrateError))
        Y1 = np.log10(np.array(integralFun.integrateError))
        Y2 = np.array(integralFun.integrateValue)
        axs[0].plot(X, Y1, linewidth=1, marker=".", label=integralFun.method)
        axs[1].plot(X, Y2, linewidth=1, marker=".", label=integralFun.method)
    axs[0].xaxis.set_major_locator(ticker.IndexLocator(base=1, offset=0))
    axs[0].yaxis.set_major_locator(ticker.MultipleLocator(0.5))
    axs[0].set_xlabel('Iteration')
    axs[0].set_ylabel('log10(error)')
    axs[0].set_title("Error-Iteration Plot")
    axs[0].legend()
    axs[1].xaxis.set_major_locator(ticker.IndexLocator(base=1, offset=0))
    axs[1].yaxis.set_major_locator(ticker.AutoLocator())
    axs[1].set_xlabel('Iteration')
    axs[1].set_ylabel('Integration Value')
    axs[1].set_title("Value-Iteration Plot")
    axs[1].legend()
    fig.savefig("1-2.png", dpi=300)
    document.add_picture("1-2.png", width=Inches(8))
    document.save('1-2.docx')
    plt.show()
